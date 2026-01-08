import io
import os
import re
from typing import List, Tuple
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, text
from models import Document, DocumentChunk
from services.ollama_service import ollama_service
from config import CHUNK_SIZE, CHUNK_OVERLAP, TOP_K_DOCUMENTS
from pypdf import PdfReader
from docx import Document as DocxDocument


class RAGService:
    """Service for RAG (Retrieval Augmented Generation)."""
    
    async def process_document(
        self,
        db: AsyncSession,
        conversation_id: str,
        filename: str,
        file_content: bytes
    ) -> str:
        """Process and store a document with embeddings."""
        
        # Determine file type
        file_ext = os.path.splitext(filename)[1].lower()
        
        # Extract text based on file type
        if file_ext == '.txt':
            text_content = file_content.decode('utf-8')
        elif file_ext == '.pdf':
            text_content = await self._extract_pdf_text(file_content)
        elif file_ext == '.docx':
            text_content = await self._extract_docx_text(file_content)
        else:
            raise ValueError(f"Unsupported file type: {file_ext}")
        
        # Create document record
        document = Document(
            conversation_id=conversation_id,
            filename=filename,
            file_type=file_ext[1:],  # Remove the dot
            content=text_content
        )
        db.add(document)
        await db.flush()  # Get document ID
        
        # Chunk the text
        chunks = self._chunk_text(text_content)
        
        # Generate embeddings and store chunks
        for idx, chunk_text in enumerate(chunks):
            embedding = await ollama_service.generate_embedding(chunk_text)
            
            chunk = DocumentChunk(
                document_id=document.id,
                chunk_text=chunk_text,
                chunk_index=idx,
                embedding=embedding
            )
            db.add(chunk)
        
        await db.commit()
        return document.id
    
    async def search_relevant_chunks(
        self,
        db: AsyncSession,
        conversation_id: str,
        query: str,
        top_k: int = TOP_K_DOCUMENTS
    ) -> List[str]:
        """Search for relevant document chunks using semantic similarity."""
        
        try:
            # Generate query embedding
            query_embedding = await ollama_service.generate_embedding(query)
            
            if not query_embedding:
                print("Warning: Failed to generate embedding for query")
                return []
            
            # Get all documents for this conversation
            result = await db.execute(
                select(Document).where(Document.conversation_id == conversation_id)
            )
            documents = result.scalars().all()
            
            if not documents:
                print(f"No documents found for conversation {conversation_id}")
                return []
            
            document_ids = [doc.id for doc in documents]
            print(f"Found {len(documents)} documents for conversation")
            
            # Search for similar chunks using pgvector
            # Using raw SQL for vector similarity search
            query_sql = text("""
                SELECT chunk_text, 1 - (embedding <=> CAST(:query_embedding AS vector)) as similarity
                FROM document_chunks
                WHERE document_id = ANY(:document_ids)
                ORDER BY embedding <=> CAST(:query_embedding AS vector)
                LIMIT :top_k
            """)
            
            result = await db.execute(
                query_sql,
                {
                    "query_embedding": str(query_embedding),  # Convert list to string
                    "document_ids": document_ids,
                    "top_k": top_k
                }
            )
            
            chunks = [row[0] for row in result.fetchall()]
            print(f"Found {len(chunks)} relevant chunks")
            return chunks
            
        except Exception as e:
            print(f"Error in search_relevant_chunks: {e}")
            import traceback
            traceback.print_exc()
            return []
    
    async def get_conversation_documents(
        self,
        db: AsyncSession,
        conversation_id: str
    ) -> List[dict]:
        """Get all documents for a conversation."""
        result = await db.execute(
            select(Document).where(Document.conversation_id == conversation_id)
        )
        documents = result.scalars().all()
        
        return [
            {
                "id": doc.id,
                "filename": doc.filename,
                "file_type": doc.file_type,
                "uploaded_at": doc.uploaded_at.isoformat()
            }
            for doc in documents
        ]
    
    async def delete_document(
        self,
        db: AsyncSession,
        document_id: str
    ) -> bool:
        """Delete a document and its chunks."""
        result = await db.execute(
            select(Document).where(Document.id == document_id)
        )
        document = result.scalars().first()
        
        if document:
            await db.delete(document)
            await db.commit()
            return True
        return False
    

    def _chunk_text(self, text: str) -> List[str]:
        """Split text into semantically meaningful chunks."""
        # Split by double newline (paragraphs)
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
                
            # Estimate size in words
            para_words = para.split()
            para_size = len(para_words)
            
            # If a single paragraph is too large, split by sentences
            if para_size > CHUNK_SIZE:
                # Flush current buffer first
                if current_chunk:
                    chunks.append("\n\n".join(current_chunk))
                    current_chunk = []
                    current_size = 0
                
                # Split large paragraph by sentences (keeping punctuation)
                sentences = re.split(r'(?<=[.!?])\s+', para)
                current_sub_chunk = []
                current_sub_size = 0
                
                for sent in sentences:
                    sent_size = len(sent.split())
                    if current_sub_size + sent_size > CHUNK_SIZE:
                        chunks.append(" ".join(current_sub_chunk))
                        current_sub_chunk = [sent]
                        current_sub_size = sent_size
                    else:
                        current_sub_chunk.append(sent)
                        current_sub_size += sent_size
                
                if current_sub_chunk:
                    chunks.append(" ".join(current_sub_chunk))
                    
            elif current_size + para_size > CHUNK_SIZE:
                # Flush buffer -> new chunk
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [para]
                current_size = para_size
            else:
                # Add to buffer
                current_chunk.append(para)
                current_size += para_size
        
        # Flush remaining buffer
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
            
        return chunks
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF."""
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX."""
        docx_file = io.BytesIO(file_content)
        doc = DocxDocument(docx_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()


# Singleton instance
rag_service = RAGService()
